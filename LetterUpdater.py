"""Author: Denzel Udemba
Date: 05/10/2023
Description: This code uses a basic word letter template and adds
data points from an Excel spreadsheet in order to make multiple
versions of the template file"""

import os
import shutil
from docx import Document
from dotenv import load_dotenv
from docx2pdf import convert
from openpyxl import load_workbook


def main():
    load_dotenv()

    # Get paths from .env file
    source_file = os.getenv("WORD_TEMPLATE_SOURCE_FILE")
    word_files_created = os.getenv("WORD_FILES_CREATED")
    pdf_files_created = os.getenv("PDF_FILES_CREATED")
    excel_file = os.getenv("XL_FILE_LOCATION")
    worksheet_name = os.getenv("WORKSHEET_NAME")

    # Load excel_dict_variables from Excel file
    excel_workbook = load_workbook(filename=excel_file, read_only=False)
    work_sheet = excel_workbook[worksheet_name]
    excel_dict_variables = {}

    for row in work_sheet.iter_rows(min_row=2, values_only=True):
        excel_key = row[0]
        excel_values = row[1:]
        excel_dict_variables[excel_key] = excel_values
        business_name = excel_dict_variables[excel_key][2]
        title = excel_dict_variables[excel_key][3]
        date = excel_dict_variables[excel_key][0]
        date_str = date.strftime('%B %d, %Y')
        date_for_title = date.strftime('%B-%d-%Y')
        contractor_name = excel_dict_variables[excel_key][1]
        target_files = [f"LightUp_ATL_GA_{file.strip()}_{date_for_title}.docx"
                        for file in excel_dict_variables[excel_key][1].split(',')]

        for new_file in target_files:
            # Copy source file to new word files directory
            shutil.copyfile(source_file, os.path.join(word_files_created, new_file))

            # Open new file for editing
            document = Document(os.path.join(word_files_created, new_file))

            # Parse through new_file & change the following field if applicable
            for paragraph in document.paragraphs:
                if "__DATE__" in paragraph.text:
                    paragraph.text = paragraph.text.replace("__DATE__", str(date_str))
                    print(f"Date: {date} changed successfully")

            for paragraph in document.paragraphs:
                if "__NAME__" in paragraph.text:
                    paragraph.text = paragraph.text.replace("__NAME__", contractor_name)
                    print(f"Contractor Name: {contractor_name} changed successfully")

            for paragraph in document.paragraphs:
                if "__NAME_OF_BUSINESS__" in paragraph.text:
                    paragraph.text = paragraph.text.replace("__NAME_OF_BUSINESS__", business_name)
                    print(f"Business Name: {business_name} changed successfully")

            for paragraph in document.paragraphs:
                if "__TITLE__" in paragraph.text:
                    paragraph.text = paragraph.text.replace("__TITLE__", title)
                    print(f"Position: {title} changed successfully")

            # Save updated file
            document.save(os.path.join(word_files_created, f"{new_file}"))
            print(f"{new_file}: saved successfully to {os.path.join(word_files_created, f'{new_file}')}")

            # Convert the file to PDF
            convert(os.path.join(word_files_created, new_file),
                    os.path.join(pdf_files_created, f"{new_file}.pdf"))

            print(f"{new_file}: converted successfully to"
                  f" {os.path.join(pdf_files_created, f'{new_file}.pdf')}")


if __name__ == '__main__':
    main()
